import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import locale
from datetime import datetime
import re

# Set the locale for formatting numbers with commas
locale.setlocale(locale.LC_ALL, '')

global formatted_text, as_on_date, promoter_name, project_name, registration_number, ECC, ICC, ECC_rehab, ICC_rehab, registration_date, \
    planning_authority, promoter_address, Diffrence_mod, today_date,ECC_FOR_REHAB,ICC_FOR_REHAB, Diffrence_FOR_REHAB,construction_cost,NEW_ECC_REHAB_5,NEW_ICC_REHAB_5,NEW_DIFFRENCE_5
formatted_text = ""  # Global variable to store the formatted text


def main():
    # Set page configuration
    st.set_page_config(page_title='Generate Form 2', page_icon='📝', layout='wide')

    # Apply custom CSS for aesthetic changes
    st.markdown(
        """
        <style>
            .stApp {
                background: linear-gradient(to right, #2b5876, #4e4376),
                            radial-gradient(circle, #F8DE22, #F94C10); /* Gradient background */

            }
            .stTitle {
                color: black !important; /* Change title color to white */
            }
            body {
                color: #000000; /* Dark gray text */
                font-size: 15px;
            }
            .sidebar {
                background-color: #7A3E65 !important; /* Background color of the sidebar */
            }
            .stFileUploader {
                background-color: #FFD93D; /* Background color inside file uploader */
                border: 2px solid white; /* Border color */
                border-radius: 10px; /* Rounded corners */
                box-shadow: 0px 4px 8px rgba(0, 0, 0, 0.8); /* Add shadow to file uploader */
                padding: 20px; /* Add padding */
            }
            .stButton>button {
                background-color: #200E3A; /* Blue button */
                color: white;
                padding: 10px 20px;
                font-size: 16px;
                border-radius: 5px;
                cursor: pointer;
            }
            .stButton>button:hover {
                background-color: #0056b3; 
            }

            .footer {
                position: fixed;
                bottom: 10px;
                left: 80%;
                transform: translateX(-50%);
                color: White; 
                font-size: 14px;
            }

        </style>
        """,
        unsafe_allow_html=True
    )

    st.title("Lets Generate Form 2 !! ✌️")

    # Sidebar
    st.sidebar.title("What you wanna Do ❔")
    selected_option = (st.sidebar.radio(" ", ["🚀 Create Form 2", "📧 Mailing Tool"]))

    if selected_option == "🚀 Create Form 2":
        create_form_2()
    elif selected_option == "📧 Mailing Tool":
        mailing_tool()

    st.markdown('<div class="footer">Created by 😎Soham</div>', unsafe_allow_html=True)


def create_form_2():
    global formatted_text, promoter_address, registration_date, planning_authority, today_date
    # File uploader widget
    uploaded_file = st.file_uploader("Dump your Form 3 here ", type=["xlsx", "xls", "xlsb"],
                                     help="Upload the new format form 3 excel.")

    # Input field for promoter address
    promoter_address = st.text_area("Enter Promoter Address", "")

    # Input field for registration date
    registration_date = st.text_input("Select Project Registration Date (dd/mm/yyyy)")

    # Input field for planning authority
    planning_authority = st.text_input("Planning Authority", "")

    # Get today's date in dd/mm/yy format
    today_date = datetime.now().strftime("%d/%m/%y")

    # Process button
    if uploaded_file is not None:
        if st.button("Process Form 3"):
            # Convert Excel to text
            text = convert_to_text(uploaded_file)
            # Process extracted text
            process_text(text)


def mailing_tool():
    st.title("Mailing Tool")
    st.write("Coming soon...")


def convert_to_text(uploaded_file):
    global formatted_text  # Access the global variable
    # Load the Excel file
    xls = pd.ExcelFile(uploaded_file)

    # Initialize an empty string to store text from all sheets
    all_text = ""

    # Iterate over each sheet and concatenate its text to all_text
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name, engine="openpyxl")
        sheet_text = df.to_string(index=False)
        all_text += sheet_text + "\n\n"

    formatted_text = all_text  # Store formatted text in the global variable
    return all_text


def process_text(text):
    global formatted_text  # Access the global variable
    # Split the text into pages
    pages = text.split("\n\n\n\n")

    # Remove unnecessary spaces and arrange page-wise
    formatted_text = ""
    for i, page in enumerate(pages):
        lines = page.strip().split("\n")  # Split the page into lines
        formatted_lines = []

        # Iterate over each line and break the line after every cell
        for line in lines:
            cells = line.split()  # Split the line into cells
            formatted_line = "|".join(cells)  # Join cells with "|" separator
            formatted_lines.append(formatted_line)

        formatted_page = "\n".join(formatted_lines)  # Join lines with newline separator
        formatted_text += f"Page {i + 1}:\n\n{formatted_page}\n\n"

    # Extract and process text variables
    global as_on_date, promoter_name, project_name, registration_number, ECC_rounded, ICC_rounded, ECC_rehab, ICC_rehab, ECC_95, ECC_5, \
        ICC_95, ICC_5, Per_1, Per_2, Difference_95, Difference_5, Diffrence_95_mod, Diffrence_5_mod, Diffrence, ECC_rehab_95, ECC_rehab_5, \
        ICC_rehab_95, ICC_rehab_5, Diffrence_rehab, Per_3, registration_date, ICC_FOR_REHAB,ECC_FOR_REHAB,Diffrence_FOR_REHAB,Diffrence_mod,construction_cost,NEW_ICC_REHAB_5,NEW_ECC_REHAB_5,NEW_DIFFRENCE_5
    


    ##CONSTANT VARIABLE
    as_on_date = extract_as_on_date().replace("|", " ")
    promoter_name = extract_promoter_name().replace("|", " ")
    project_name = extract_project_name().replace("|", " ")
    registration_number = extract_registration_number().replace("|", "")
    ECC = extract_ECC().replace("|", " ")
    ICC = extract_ICC().replace("|", " ")
    ECC_rehab = extract_ECC_rehab().replace("|", " ")
    ICC_rehab = extract_ICC_rehab().replace("|", " ")
    ## VARIABLE FOR NORMAL CASE
    # Calculate ECC_95, ECC_5, ICC_95, and ICC_5
    ECC_rounded = round(float(ECC))
    ICC_rounded = round(float(ICC))
    ECC_rehab_rounded = round(float(ECC_rehab))
    ICC_rehab_rounded = round(float(ICC_rehab))
    ECC_95 = round(0.95 * ECC_rounded)
    ECC_5 = round(0.05 * ECC_rounded)
    ICC_95 = round(0.95 * ICC_rounded)
    ICC_5 = round(0.05 * ICC_rounded)
    ECC_rehab_95 = round(0.95 * ECC_rehab_rounded)
    ECC_rehab_5 = round(0.05 * ECC_rehab_rounded)
    ICC_rehab_95 = round(0.95 * ICC_rehab_rounded)
    ICC_rehab_5 = round(0.05 * ICC_rehab_rounded)

    ## VARIABLE FOR REHAB CASE
    ECC_FOR_REHAB = ECC_rounded + ECC_rehab_rounded
    ICC_FOR_REHAB = ICC_rounded + ICC_rehab_rounded
    NEW_ECC_REHAB_5 = (ECC_FOR_REHAB * 0.05)
    NEW_ICC_REHAB_5 = ICC_FOR_REHAB * 0.05)
    Diffrence_FOR_REHAB = ECC_FOR_REHAB - ICC_FOR_REHAB

    # Workdone percentage and difference calculation
    Per_1 = float(round((ICC_95 / ECC_95) * 100, 2))
    Per_2 = float(round((ICC_5 / ECC_5) * 100, 2))
    if ECC_rehab and ICC_rehab and float(ECC_rehab) != 0:
        Per_3 = float(round((ICC_rehab_95 / ECC_rehab_95) * 100, 2))
    else:
        Per_3 = None
    Difference_95 = ECC_95 - ICC_95
    Difference_5 = ECC_5 - ICC_5
    Diffrence_95_mod = abs(Difference_95)
    Diffrence_5_mod = abs(ECC_5 - ICC_5)
    Diffrence = round(float(ECC_rounded - ICC_rounded))
    Diffrence_rehab = (ECC_rehab_95 - ICC_rehab_95)
    Diffrence_mod = abs(ECC_rounded - ICC_rounded)
    construction_cost = Diffrence_95_mod - Diffrence_5_mod
    NEW_DIFFRENCE_5 = NEW_ECC_REHAB_5 - NEW_ICC_REHAB_5

    # Display extracted values
    st.subheader("Values Extracted From The Given Excel!:")
    st.write(f"As on date: {as_on_date}")
    st.write(f"Promoter name: {promoter_name}")
    st.write(f"Project name: {project_name}")
    st.write(f"RERA number: {registration_number}")
    st.write(f"Estimated Construction Cost: {ECC}")
    st.write(f"Incurred Construction Cost: {ICC}")
    st.write(f"Modal 5: {Diffrence_5_mod}")
    st.write(f"Modal 95: {Diffrence_95_mod}")
    st.write(f"{Diffrence_mod}")

    # Edit the Word document and offer download
    edited_docx_bytes = edit_docx(as_on_date, promoter_name, ECC_95)
    st.download_button(label="Download Edited Document", data=edited_docx_bytes,
                       file_name="Machine_generated_form_2.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

def extract_as_on_date():
    global formatted_text  # Access the global variable
    as_on_date = formatted_text.split("As|on|")[1].split("NaN|NaN")[0].strip()
    return as_on_date

def extract_promoter_name():
    global formatted_text  # Access the global variable
    promoter_name = formatted_text.split("|being|developed|by|")[1].split("|NaN|NaN")[0].strip()
    return promoter_name

def extract_project_name():
    global formatted_text  # Access the global variable
    project_name = \
    formatted_text.split("This|certificate|is|being|issued|for|the|")[1].split("|having|MahaRERA|Registration|")[
        0].strip()
    return project_name

def extract_registration_number():
    global formatted_text  # Access the global variable
    registration_number = formatted_text.split("|MahaRERA|Registration|Number|")[1].split("|being|developed|")[
        0].strip()
    return registration_number

def extract_ECC():
    global formatted_text  # Access the global variable
    ECC = formatted_text.split("a|Estimated|Cost|of|Construction|as|certified|by|Engineer.|")[1].split("b.|")[0].strip()
    return ECC

def extract_ICC():
    global formatted_text  # Access the global variable
    ICC = formatted_text.split(
        "(b)|Actual|Cost|of|construction|incurred|as|per|the|books|of|accounts|as|verified|by|the|CA.|")[1].split(
        "(ii)|")[0].strip()
    return ICC

def extract_ECC_rehab():
    global formatted_text  # Access the global variable
    ECC_rehab = formatted_text.split(
        "|Estimated|Construction|Cost|of|Rehab|Building|including|Site|Development|and|Infrastructure|for|the|same|as|certified|by|the|Engineer.|")[
        1].split("(ii)|")[0].strip()
    return ECC_rehab if ECC_rehab != "NaN" else "0"


def extract_ICC_rehab():
    global formatted_text  # Access the global variable
    ICC_rehab = formatted_text.split(
        "|Incurred|Expenditure|for|construction|Rehab|building|as|per|the|books|of|accounts|as|verified|by|the|CA.|")[
        1].split("(ii)|")[0].strip()
    return ICC_rehab if ICC_rehab != "NaN" else "0"

def format_for_float(number):
    number_float = str(number)
    return number_float
# Function to format numbers with commas
def format_number_with_commas(number):
    # Convert number to string
    number_str = str(int(number))  # Convert to int to remove decimal point if present

    # Determine the length of the number string
    length = len(number_str)

    # Apply formatting based on the number of digits
    if length == 11:
        formatted_number = number_str[:-9] + "," + number_str[-9:-7] + "," + number_str[-7:-5] + "," + number_str[-5:-3] + "," + number_str[-3:]
    elif length == 10:
        formatted_number = number_str[:-9] + "," + number_str[-9:-7] + "," + number_str[-7:-5] + "," + number_str[-5:-3] + "," + number_str[-3:]
    elif length == 9:
        formatted_number = number_str[:-7] + "," + number_str[-7:-5] + "," + number_str[-5:-3] + "," + number_str[-3:]
    elif length == 8:
        formatted_number = number_str[:-7] + "," + number_str[-7:-5] + "," + number_str[-5:-3] + "," + number_str[-3:]
    elif length == 7:
        formatted_number = number_str[:-5] + "," + number_str[-5:-3] + "," + number_str[-3:]
    elif length == 6:
        formatted_number = number_str[:-5] + "," + number_str[-5:-3] + "," + number_str[-3:]
    elif length == 5:
        formatted_number = number_str[:-3] + "," + number_str[-3:]
    elif length == 4:
        formatted_number = number_str[:-3] + "," + number_str[-3:]
    elif length == 3:
        formatted_number = number_str
    elif length == 2:
        formatted_number = number_str
    else:
        formatted_number = number_str

    return formatted_number


def edit_docx(as_on_date, promoter_name, ECC_95):
    # Check if ECC_rehab and ICC_rehab are not null or zero
    if (ECC_rehab != "" and ICC_rehab != "") or (ECC_rehab != "" and ICC_rehab == "") or (ECC_rehab == "" and ICC_rehab != ""):
        # Use the template for cases where ECC_rehab and ICC_rehab are not null or zero
        template_path = "form_2_rehab.docx"
    elif Diffrence < 0:
        # Use the template for cases where Diffrence is negative
        template_path = "form_2_exceptional.docx"
    else:
        template_path = "form_2_normal.docx"

    # Load the Word template based on the determined path
    doc = Document(template_path)

    # Define font settings
    font_name = "Gadugi"
    font_size = Pt(12)  # Adjust font size as needed

    # Replace placeholders in table cells with variable values and set font settings
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "ECC_95" in cell.text:
                    cell.text = cell.text.replace("ECC_95", format_number_with_commas(ECC_95))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "ECC_5" in cell.text:
                    cell.text = cell.text.replace("ECC_5", format_number_with_commas(ECC_5))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "ICC_95" in cell.text:
                    cell.text = cell.text.replace("ICC_95", format_number_with_commas(ICC_95))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "ICC_5" in cell.text:
                    cell.text = cell.text.replace("ICC_5", format_number_with_commas(ICC_5))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "Per_1" in cell.text:
                    cell.text = cell.text.replace("Per_1", format_for_float(Per_1))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "Per_2" in cell.text:
                    cell.text = cell.text.replace("Per_2", format_for_float(Per_2))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "Diffrence_95" in cell.text:
                    cell.text = cell.text.replace("Diffrence_95", format_number_with_commas(Difference_95))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "Diffrence_5" in cell.text:
                    cell.text = cell.text.replace("Diffrence_5", format_number_with_commas(Difference_5))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "as date" in cell.text:
                    cell.text = cell.text.replace("as date", as_on_date)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                            if "as date" in run.text:
                                run.bold = True
                if "reg_date" in cell.text:
                    cell.text = cell.text.replace("reg_date", registration_date)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                            if "as date" in run.text:
                                run.bold = True
                if "ECC_rehab_95" in cell.text:
                    cell.text = cell.text.replace("ECC_rehab_95", format_number_with_commas(ECC_rehab_95))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "ICC_rehab_95" in cell.text:
                    cell.text = cell.text.replace("ICC_rehab_95", format_number_with_commas(ICC_rehab_95))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "Per_3" in cell.text:
                    cell.text = cell.text.replace("Per_3", format_for_float(Per_3))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "Diffrence_rehab" in cell.text:
                    cell.text = cell.text.replace("Diffrence_rehab", format_number_with_commas(Diffrence_rehab))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "NEW_ICC_5" in cell.text:
                    cell.text = cell.text.replace("NEW_ICC_5", format_number_with_commas(NEW_ICC_REHAB_5))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "NEW_ECC_5" in cell.text:
                    cell.text = cell.text.replace("NEW_ECC_5", format_number_with_commas(NEW_ECC_REHAB_5))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "diff_95_mod" in cell.text:
                    cell.text = cell.text.replace("diff_95_mod", format_number_with_commas(Diffrence_95_mod))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "diff_5_mod" in cell.text:
                    cell.text = cell.text.replace("diff_5_mod", format_number_with_commas(Diffrence_5_mod))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "ccc" in cell.text:
                    cell.text = cell.text.replace("ccc", format_number_with_commas(construction_cost))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size

            # Replace placeholders in text paragraphs with variable values and set font settings
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "as date" in run.text:
                run.text = run.text.replace("as date", as_on_date)
                run.font.name = font_name
                run.font.size = font_size
            elif "promoter_name" in run.text:
                run.text = run.text.replace("promoter_name", promoter_name)
                run.font.name = font_name
                run.font.size = font_size
            elif "project_name" in run.text:
                run.text = run.text.replace("project_name", project_name)
                run.font.name = font_name
                run.font.size = font_size
            elif "RERA_number" in run.text:
                run.text = run.text.replace("RERA_number", registration_number)
                run.font.name = font_name
                run.font.size = font_size
            elif "ECC" in run.text:
                run.text = run.text.replace("ECC", format_number_with_commas(ECC_rounded))
                run.font.name = font_name
                run.font.size = font_size
            elif "ICC" in run.text:
                run.text = run.text.replace("ICC", format_number_with_commas(ICC_rounded))
                run.font.name = font_name
                run.font.size = font_size
            elif "Diffrence" in run.text:
                run.text = run.text.replace("Diffrence", format_number_with_commas(Diffrence))
                run.font.name = font_name
                run.font.size = font_size
            elif "planning_authority_name" in run.text:
                run.text = run.text.replace("planning_authority_name", planning_authority)
                run.font.name = font_name
                run.font.size = font_size
            elif "promoter_address" in run.text:
                run.text = run.text.replace("promoter_address", promoter_address)
                run.font.name = font_name
                run.font.size = font_size
            elif "reg_date" in cell.text:
                cell.text = cell.text.replace("reg_date", registration_date.strftime("%d/%m/%Y"))
                run.font.name = font_name
                run.font.size = font_size
            elif "modulus_difference" in run.text:
                run.text = run.text.replace("modulus_difference", format_number_with_commas(Diffrence_mod))
                run.font.name = font_name
                run.font.size = font_size

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "for_rehab_ecc" in run.text:
                run.text = run.text.replace("for_rehab_ecc", format_number_with_commas(ECC_FOR_REHAB))
                run.font.name = font_name
                run.font.size = font_size
            if "for_rehab_icc" in run.text:
                run.text = run.text.replace("for_rehab_icc", format_number_with_commas(ICC_FOR_REHAB))
                run.font.name = font_name
                run.font.size = font_size
            if "for_rehab_diff" in run.text:
                run.text = run.text.replace("for_rehab_diff", format_number_with_commas(Diffrence_FOR_REHAB))
                run.font.name = font_name
                run.font.size = font_size

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "today_date" in run.text:
                run.text = run.text.replace("today_date", today_date)
                run.font.name = font_name
                run.font.size = font_size


    # Save the edited document to a BytesIO object
    edited_docx_bytes = BytesIO()
    doc.save(edited_docx_bytes)
    edited_docx_bytes.seek(0)  # Reset the file pointer to the beginning
    return edited_docx_bytes

if __name__ == "__main__":
    main()
