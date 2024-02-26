import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import locale

# Set the locale for formatting numbers with commas
locale.setlocale(locale.LC_ALL, '')

global formatted_text
global as_on_date
global promoter_name
global project_name
global registration_number
global ECC
global ICC
global ECC_rehab
global ICC_rehab
global registration_date
global planning_authority
global promoter_address
global Diffrence_mod

formatted_text = ""  # Global variable to store the formatted text

def main():
    st.markdown(
        """
        <link rel="stylesheet" href="styles.css">
        """,
        unsafe_allow_html=True
    )
    global formatted_text , promoter_address, registration_date,planning_authority
    st.title("Excel to Text Converter")

    # File uploader widget
    uploaded_file = st.file_uploader("Upload an Excel file", type=["xlsx", "xls", "xlsb"])

    # Input field for promoter address
    promoter_address = st.text_area("Enter Promoter Address (Multi-line)", "")

    # Input field for registration date
    registration_date = st.date_input("Enter Registration Date (dd/mm/yyyy)")

    # Input field for planning authority
    planning_authority = st.text_input("Enter Planning Authority", "")

    # Process button
    if uploaded_file is not None:
        if st.button("Convert and Process"):
            # Convert Excel to text
            text = convert_to_text(uploaded_file)

            # Process extracted text
            process_text(text)

def convert_to_text(uploaded_file):
    global formatted_text  # Access the global variable
    # Load the Excel file
    xls = pd.ExcelFile(uploaded_file)

    # Initialize an empty string to store text from all sheets
    all_text = ""

    # Iterate over each sheet and concatenate its text to all_text
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name, engine="openpyxl")
        sheet_text = df.to_string(index=False)
        all_text += sheet_text + "\n\n"

    formatted_text = all_text  # Store formatted text in the global variable
    return all_text

def process_text(text):
    global formatted_text  # Access the global variable
    # Split the text into pages
    pages = text.split("\n\n\n\n")

    # Remove unnecessary spaces and arrange page-wise
    formatted_text = ""
    for i, page in enumerate(pages):
        lines = page.strip().split("\n")  # Split the page into lines
        formatted_lines = []

        # Iterate over each line and break the line after every cell
        for line in lines:
            cells = line.split()  # Split the line into cells
            formatted_line = "|".join(cells)  # Join cells with "|" separator
            formatted_lines.append(formatted_line)

        formatted_page = "\n".join(formatted_lines)  # Join lines with newline separator
        formatted_text += f"Page {i+1}:\n\n{formatted_page}\n\n"

    # Extract and process text variables
    global as_on_date
    global promoter_name
    global project_name
    global registration_number
    global ECC_rounded
    global ICC_rounded
    global ECC_rehab
    global ICC_rehab
    global ECC_95
    global ECC_5
    global ICC_95
    global ICC_5
    global Per_1
    global Per_2
    global Difference_95
    global Difference_5
    global Diffrence_95_mod
    global Diffrence_5_mod
    global Diffrence
    global ECC_rehab_95
    global ECC_rehab_5
    global ICC_rehab_95
    global ICC_rehab_5
    global Diffrence_rehab
    global Per_3

    as_on_date = extract_as_on_date().replace("|", " ")
    promoter_name = extract_promoter_name().replace("|", " ")
    project_name = extract_project_name().replace("|", " ")
    registration_number = extract_registration_number().replace("|", "")
    ECC = extract_ECC().replace("|", " ")
    ICC = extract_ICC().replace("|", " ")
    ECC_rehab = extract_ECC_rehab().replace("|", " ")
    ICC_rehab = extract_ICC_rehab().replace("|", " ")

    # Calculate ECC_95, ECC_5, ICC_95, and ICC_5
    ECC_rounded = round(float(ECC))
    ICC_rounded = round(float(ICC))
    ECC_rehab_rounded = round(float(ECC_rehab))
    ICC_rehab_rounded = round(float(ICC_rehab))
    ECC_95 = round(0.95 * ECC_rounded)
    ECC_5 = round(0.05 * ECC_rounded)
    ICC_95 = round(0.95 * ICC_rounded)
    ICC_5 = round(0.05 * ICC_rounded)
    ECC_rehab_95 = round(0.95 * ECC_rehab_rounded)
    ECC_rehab_5 = round(0.05 * ECC_rehab_rounded)
    ICC_rehab_95 = round(0.95 * ICC_rehab_rounded)
    ICC_rehab_5 = round(0.05 * ICC_rehab_rounded)


    # Workdone percentage and difference calculation
    Per_1 = round((ICC_95/ECC_95) * 100,2)
    Per_2 = round((ICC_5/ECC_5) * 100,2)
    if ECC_rehab and ICC_rehab and float(ECC_rehab) != 0:
        Per_3 = round((ICC_rehab_95 / ECC_rehab_95) * 100, 2)
    else:
        Per_3 = None
    Difference_95 = ECC_95 - ICC_95
    Difference_5 = ECC_5 - ICC_5
    Diffrence_95_mod = abs(Difference_95)
    Diffrence_5_mod = abs(ECC_5 - ICC_5)
    Diffrence = round(float(ECC_rounded - ICC_rounded))
    Diffrence_rehab = (ECC_rehab_95 - ICC_rehab_95)

    # Display the formatted text
    st.subheader("Extracted Text:")
    st.text(formatted_text)

    # Display extracted values
    st.subheader("Extracted Values:")
    st.write(f"As on date: {as_on_date}")
    st.write(f"Promoter name: {promoter_name}")
    st.write(f"Project name: {project_name}")
    st.write(f"RERA number: {registration_number}")
    st.write(f"Estimated Construction Cost: {ECC}")
    st.write(f"Incurred Construction Cost: {ICC}")
    st.write(f"ECC rehab: {ECC_rehab}")
    st.write(f"ICC rehab: {ICC_rehab}")

    # Display calculated values
    st.subheader("Calculated Values:")
    st.write(f"ECC (rounded): {format_number_with_commas(ECC_rounded)}")
    st.write(f"ECC_95: {format_number_with_commas(ECC_95)}")
    st.write(f"ECC_5: {format_number_with_commas(ECC_5)}")
    st.write(f"ICC (rounded): {format_number_with_commas(ICC_rounded)}")
    st.write(f"ICC_95: {format_number_with_commas(ICC_95)}")
    st.write(f"ICC_5: {format_number_with_commas(ICC_5)}")
    st.write(f"Per_1: {format_number_with_commas(Per_1)}")
    st.write(f"Per_2: {format_number_with_commas(Per_2)}")
    st.write(f"Difference_95: {format_number_with_commas(Difference_95)}")
    st.write(f"Difference_5: {format_number_with_commas(Difference_5)}")
    st.write(f"{ECC_rehab_5}, {ECC_rehab_95}")
    st.write(f"{ICC_rehab_5}, {ICC_rehab_95}")

    # Edit the Word document and offer download
    edited_docx_bytes = edit_docx(as_on_date, promoter_name, ECC_95)
    st.download_button(label="Download Edited Document", data=edited_docx_bytes, file_name="edited_document.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

def extract_as_on_date():
    global formatted_text  # Access the global variable
    as_on_date = formatted_text.split("As|on|")[1].split("NaN|NaN")[0].strip()
    return as_on_date

def extract_promoter_name():
    global formatted_text  # Access the global variable
    promoter_name = formatted_text.split("|being|developed|by|")[1].split("|NaN|NaN")[0].strip()
    return promoter_name

def extract_project_name():
    global formatted_text  # Access the global variable
    project_name = formatted_text.split("This|certificate|is|being|issued|for|the|")[1].split("|having|MahaRERA|Registration|")[0].strip()
    return project_name

def extract_registration_number():
    global formatted_text  # Access the global variable
    registration_number = formatted_text.split("|MahaRERA|Registration|Number|")[1].split("|being|developed|")[0].strip()
    return registration_number

def extract_ECC():
    global formatted_text  # Access the global variable
    ECC = formatted_text.split("a|Estimated|Cost|of|Construction|as|certified|by|Engineer.|")[1].split("b.|")[0].strip()
    return ECC

def extract_ICC():
    global formatted_text  # Access the global variable
    ICC = formatted_text.split("(b)|Actual|Cost|of|construction|incurred|as|per|the|books|of|accounts|as|verified|by|the|CA.|")[1].split("(ii)|")[0].strip()
    return ICC

def extract_ECC_rehab():
    global formatted_text  # Access the global variable
    ECC_rehab = formatted_text.split("|Estimated|Construction|Cost|of|Rehab|Building|including|Site|Development|and|Infrastructure|for|the|same|as|certified|by|the|Engineer.|")[1].split("(ii)|")[0].strip()
    return ECC_rehab

def extract_ICC_rehab():
    global formatted_text  # Access the global variable
    ICC_rehab = formatted_text.split("|Incurred|Expenditure|for|construction|Rehab|building|as|per|the|books|of|accounts|as|verified|by|the|CA.|")[1].split("(ii)|")[0].strip()
    return ICC_rehab

# Function to format numbers with commas
def format_number_with_commas(number):
    return locale.format_string("%d", number, grouping=True)

def edit_docx(as_on_date, promoter_name, ECC_95):
        # Check if ECC_rehab and ICC_rehab are not null or zero
    if ECC_rehab != "" and ICC_rehab != "" and float(ECC_rehab) != 0 and float(ICC_rehab) != 0:
            # Use the template for cases where ECC_rehab and ICC_rehab are not null or zero
            template_path = "form_2_rehab.docx"
    elif Diffrence < 0:
        # Use the template for cases where Diffrence is negative
        template_path = "form_2_exceptional.docx"
    else:
        template_path = "form_2_normal.docx"

        # Load the Word template based on the determined path
    doc = Document(template_path)

    # Define font settings
    font_name = "Gadugi"
    font_size = Pt(12)  # Adjust font size as needed

    # Replace placeholders in table cells with variable values and set font settings
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "ECC_95" in cell.text:
                    cell.text = cell.text.replace("ECC_95", format_number_with_commas(ECC_95))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "ECC_5" in cell.text:
                    cell.text = cell.text.replace("ECC_5", format_number_with_commas(ECC_5))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "ICC_95" in cell.text:
                    cell.text = cell.text.replace("ICC_95", format_number_with_commas(ICC_95))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "ICC_5" in cell.text:
                    cell.text = cell.text.replace("ICC_5", format_number_with_commas(ICC_5))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "Per_1" in cell.text:
                    cell.text = cell.text.replace("Per_1", format_number_with_commas(Per_1))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "Per_2" in cell.text:
                    cell.text = cell.text.replace("Per_2", format_number_with_commas(Per_2))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "Diffrence_95" in cell.text:
                    cell.text = cell.text.replace("Diffrence_95", format_number_with_commas(Difference_95))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "Diffrence_5" in cell.text:
                    cell.text = cell.text.replace("Diffrence_5", format_number_with_commas(Difference_5))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "as date" in cell.text:
                    cell.text = cell.text.replace("as date", as_on_date)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                            if "as date" in run.text:
                                run.bold = True  # Make text bold
                if "ECC_rehab_95" in cell.text:
                    cell.text = cell.text.replace("ECC_rehab_95", format_number_with_commas(ECC_rehab_95))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "ICC_rehab_95" in cell.text:
                    cell.text = cell.text.replace("ICC_rehab_95", format_number_with_commas(ICC_rehab_95))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "Per_3" in cell.text:
                    cell.text = cell.text.replace("Per_3", format_number_with_commas(Per_3))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True  # Make text bold
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "Diffrence_rehab" in cell.text:
                    cell.text = cell.text.replace("Diffrence_rehab", format_number_with_commas(Diffrence_rehab))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "Diffrence_95_mod" in cell.text:
                    cell.text = cell.text.replace("Diffrence_95_mod", format_number_with_commas(Diffrence_95_mod))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size
                if "Diffrence_5_mod" in cell.text:
                    cell.text = cell.text.replace("Diffrence_5_mod", format_number_with_commas(Diffrence_5_mod))
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Justify to center
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = font_name  # Set font name
                            run.font.size = font_size

    # Replace placeholders in text paragraphs with variable values and set font settings
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "as date" in run.text:
                run.text = run.text.replace("as date", as_on_date)
                run.font.name = font_name
                run.font.size = font_size
            if "promoter_name" in run.text:
                run.text = run.text.replace("promoter_name", promoter_name)
                run.font.name = font_name
                run.font.size = font_size
            if "project_name" in run.text:
                run.text = run.text.replace("project_name", project_name)
                run.font.name = font_name
                run.font.size = font_size
            if "RERA_number" in run.text:
                run.text = run.text.replace("RERA_number", registration_number)
                run.font.name = font_name
                run.font.size = font_size
            if "ECC" in run.text:
                run.text = run.text.replace("ECC", format_number_with_commas(ECC_rounded))
                run.font.name = font_name
                run.font.size = font_size
            if "ICC" in run.text:
                run.text = run.text.replace("ICC", format_number_with_commas(ICC_rounded))
                run.font.name = font_name
                run.font.size = font_size
            if "Diffrence" in run.text:
                run.text = run.text.replace("Diffrence", format_number_with_commas(Diffrence))
                run.font.name = font_name
                run.font.size = font_size

            if "planning_authority_name" in run.text:
                run.text = run.text.replace("planning_authority_name", planning_authority)
                run.font.name = font_name
                run.font.size = font_size
            if "promoter_address" in run.text:
                run.text = run.text.replace("promoter_address", promoter_address)
                run.font.name = font_name
                run.font.size = font_size
            if "reg_date" in run.text:
                run.text = run.text.replace("reg_date", registration_date)
                run.font.name = font_name
                run.font.size = font_size
            if "modulus_diffrence" in run.text:
                run.text = run.text.replace("modulus_diffrence", format_number_with_commas(Diffrence_mod))
                run.font.name = font_name
                run.font.size = font_size
            # Add more conditions for other placeholders if needed

    # Save the modified Word document to a BytesIO object
    edited_docx_bytes = BytesIO()
    doc.save(edited_docx_bytes)

    return edited_docx_bytes.getvalue()


if __name__ == "__main__":
    main()
